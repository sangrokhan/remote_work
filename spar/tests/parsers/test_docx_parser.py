from __future__ import annotations

from pathlib import Path

import pytest
import docx as _docx

from spar.parsers.docx_config import DocxParseConfig


class TestDocxParseConfig:
    def test_defaults(self) -> None:
        cfg = DocxParseConfig()
        assert cfg.heading_depth == 2
        assert cfg.output_dir == Path("output")
        assert cfg.slugify_max_len == 30

    def test_custom_values(self) -> None:
        cfg = DocxParseConfig(heading_depth=3, output_dir=Path("/tmp/out"), slugify_max_len=20)
        assert cfg.heading_depth == 3
        assert cfg.output_dir == Path("/tmp/out")
        assert cfg.slugify_max_len == 20


from spar.parsers.docx_parser import _slugify


class TestSlugify:
    def test_spaces_to_dash(self) -> None:
        assert _slugify("System Overview", 30) == "System-Overview"

    def test_truncates(self) -> None:
        result = _slugify("A Very Long Section Title Here", 15)
        assert len(result) <= 15

    def test_strips_special_chars(self) -> None:
        assert _slugify("Section (v2.0)!", 30) == "Section-v2.0"

    def test_empty_string(self) -> None:
        assert _slugify("", 30) == "unnamed"


from spar.parsers.docx_parser import ParseResult, ExtractedTable, ExtractedImage


class TestParseResult:
    def test_parse_result_fields(self, tmp_path: Path) -> None:
        table = ExtractedTable(
            path=tmp_path / "Table_Intro_1.csv",
            section_path=["Introduction"],
            seq=1,
        )
        image = ExtractedImage(
            path=tmp_path / "Fig_Intro_1.png",
            section_path=["Introduction"],
            seq=1,
            ext="png",
        )
        result = ParseResult(markdown="# Hello", tables=[table], images=[image])
        assert result.markdown == "# Hello"
        assert len(result.tables) == 1
        assert len(result.images) == 1
        assert result.tables[0].seq == 1
        assert result.images[0].ext == "png"


import io
from spar.parsers.docx_parser import DocxParser


def _make_docx(tmp_path, setup_fn):
    import docx as _docx
    doc = _docx.Document()
    setup_fn(doc)
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path


class TestDocxParserHeadings:
    def test_heading1_becomes_h1(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Introduction", level=1)
            doc.add_paragraph("Some text here.")

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        parser = DocxParser(cfg)
        result = parser.parse(path)
        assert "# Introduction" in result.markdown
        assert "Some text here." in result.markdown

    def test_heading2_within_depth_becomes_h2(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Chapter", level=1)
            doc.add_heading("Sub Section", level=2)
            doc.add_paragraph("Body text.")

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(heading_depth=2, output_dir=tmp_path)
        parser = DocxParser(cfg)
        result = parser.parse(path)
        assert "# Chapter" in result.markdown
        assert "## Sub Section" in result.markdown

    def test_heading_beyond_depth_still_rendered(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Chapter", level=1)
            doc.add_heading("Deep", level=3)

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(heading_depth=2, output_dir=tmp_path)
        parser = DocxParser(cfg)
        result = parser.parse(path)
        assert "### Deep" in result.markdown


class TestDocxParserTables:
    def test_table_csv_created(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Config", level=1)
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Param"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "TTT"
            table.cell(1, 1).text = "100ms"

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)
        assert len(result.tables) == 1
        csv_path = result.tables[0].path
        assert csv_path.exists()
        assert "Table_Config_1" in csv_path.name

    def test_table_csv_content(self, tmp_path) -> None:
        import csv
        def setup(doc):
            doc.add_heading("Parameters", level=1)
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Alpha"
            table.cell(1, 1).text = "1.0"

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)
        csv_path = result.tables[0].path
        lines = csv_path.read_text(encoding="utf-8").splitlines()
        assert lines[0].startswith("# section:")
        assert lines[1].startswith("# source:")
        data_lines = [l for l in lines if not l.startswith("#")]
        rows = list(csv.reader(data_lines))
        assert rows[0] == ["Name", "Value"]
        assert rows[1] == ["Alpha", "1.0"]

    def test_table_placeholder_in_markdown(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Intro", level=1)
            table = doc.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "X"

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)
        assert "<!-- TABLE: Table_Intro_1 -->" in result.markdown

    def test_table_seq_resets_per_section(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Sec A", level=1)
            t1 = doc.add_table(rows=1, cols=1)
            t1.cell(0, 0).text = "A1"
            doc.add_heading("Sec B", level=1)
            t2 = doc.add_table(rows=1, cols=1)
            t2.cell(0, 0).text = "B1"

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)
        assert result.tables[0].seq == 1
        assert result.tables[1].seq == 1
        assert "Sec-A" in result.tables[0].path.name
        assert "Sec-B" in result.tables[1].path.name


class TestDocxParserImages:
    # Minimal valid 1x1 PNG bytes (no Pillow needed)
    _MINIMAL_PNG = bytes([
        0x89,0x50,0x4E,0x47,0x0D,0x0A,0x1A,0x0A,0x00,0x00,0x00,0x0D,0x49,0x48,0x44,0x52,
        0x00,0x00,0x00,0x01,0x00,0x00,0x00,0x01,0x08,0x02,0x00,0x00,0x00,0x90,0x77,0x53,
        0xDE,0x00,0x00,0x00,0x0C,0x49,0x44,0x41,0x54,0x08,0xD7,0x63,0xF8,0xCF,0xC0,0x00,
        0x00,0x00,0x02,0x00,0x01,0xE2,0x21,0xBC,0x33,0x00,0x00,0x00,0x00,0x49,0x45,0x4E,
        0x44,0xAE,0x42,0x60,0x82,
    ])

    def test_image_file_created(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Diagrams", level=1)
            doc.add_picture(io.BytesIO(self._MINIMAL_PNG))

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)

        assert len(result.images) == 1
        img_path = result.images[0].path
        assert img_path.exists()
        assert "Fig_Diagrams_1" in img_path.name

    def test_image_meta_file_created(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Diagrams", level=1)
            doc.add_picture(io.BytesIO(self._MINIMAL_PNG))

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)

        meta_path = Path(str(result.images[0].path) + ".meta")
        assert meta_path.exists()
        content = meta_path.read_text(encoding="utf-8")
        assert "section: Diagrams" in content
        assert "seq: 1" in content

    def test_image_placeholder_in_markdown(self, tmp_path) -> None:
        def setup(doc):
            doc.add_heading("Diagrams", level=1)
            doc.add_picture(io.BytesIO(self._MINIMAL_PNG))

        path = _make_docx(tmp_path, setup)
        cfg = DocxParseConfig(output_dir=tmp_path)
        result = DocxParser(cfg).parse(path)
        assert "<!-- IMAGE: Fig_Diagrams_1" in result.markdown
