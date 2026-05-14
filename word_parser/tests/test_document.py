import io
import pytest
from docx import Document as DocxDocument
from docx.shared import Pt
from tests.conftest import (
    make_docx, add_heading, add_paragraph, add_paragraph_with_font_size,
    add_page_break, add_table, docx_bytes,
)
from core.document import stream_elements
from core.models import ParagraphElement, TableElement, ImageElement


def test_stream_paragraphs():
    doc = make_docx()
    add_paragraph(doc, "Hello world")
    elements = list(stream_elements(docx_bytes(doc)))
    paras = [e for e in elements if isinstance(e, ParagraphElement)]
    texts = [p.text for p in paras]
    assert "Hello world" in texts


def test_stream_heading_style():
    doc = make_docx()
    add_heading(doc, "Section 1", level=1)
    elements = list(stream_elements(docx_bytes(doc)))
    paras = [e for e in elements if isinstance(e, ParagraphElement)]
    heading = next(p for p in paras if p.text == "Section 1")
    assert heading.style_name == "Heading 1"


def test_stream_table():
    doc = make_docx()
    add_table(doc, [["Name", "Value"], ["foo", "bar"]])
    elements = list(stream_elements(docx_bytes(doc)))
    tables = [e for e in elements if isinstance(e, TableElement)]
    assert len(tables) == 1
    assert tables[0].col_count == 2
    assert tables[0].rows[0] == ["Name", "Value"]


def test_stream_page_break_paragraph():
    doc = make_docx()
    add_paragraph(doc, "Before")
    add_page_break(doc)
    add_paragraph(doc, "After")
    elements = list(stream_elements(docx_bytes(doc)))
    paras = [e for e in elements if isinstance(e, ParagraphElement)]
    page_breaks = [p for p in paras if p.is_page_break]
    assert len(page_breaks) >= 1


def test_stream_table_preceded_by_page_break():
    doc = make_docx()
    add_table(doc, [["A"], ["1"]])
    add_page_break(doc)
    add_table(doc, [["A"], ["2"]])
    elements = list(stream_elements(docx_bytes(doc)))
    tables = [e for e in elements if isinstance(e, TableElement)]
    assert len(tables) == 2
    assert tables[1].preceded_by_page_break is True


def test_stream_font_size_on_run():
    doc = make_docx()
    add_paragraph_with_font_size(doc, "Big text", 24.0)
    elements = list(stream_elements(docx_bytes(doc)))
    paras = [e for e in elements if isinstance(e, ParagraphElement)]
    big = next(p for p in paras if p.text == "Big text")
    assert any(r.font_size == 24.0 for r in big.runs)
