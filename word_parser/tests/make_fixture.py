"""Generate sample.docx test fixture with headings, tables, and images."""
import io
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


def _set_vmerge(tc, restart: bool) -> None:
    """Add proper vMerge element to a cell's tcPr. restart=True → start of span, False → continuation."""
    from lxml import etree
    tc_pr = tc.find(qn("w:tcPr"))
    if tc_pr is None:
        tc_pr = etree.SubElement(tc, qn("w:tcPr"))
        tc.insert(0, tc_pr)
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is not None:
        tc_pr.remove(v_merge)
    v_merge = etree.SubElement(tc_pr, qn("w:vMerge"))
    if restart:
        v_merge.set(qn("w:val"), "restart")


def _png_bytes(color: tuple, width: int = 200, height: int = 120, label: str = "") -> bytes:
    img = Image.new("RGB", (width, height), color)
    draw = ImageDraw.Draw(img)
    if label:
        draw.text((10, height // 2 - 8), label, fill=(255, 255, 255))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def build() -> Document:
    doc = Document()

    # ── H1: 소개 ────────────────────────────────────────────────────────────
    doc.add_heading("1. 프로젝트 개요", level=1)
    doc.add_paragraph(
        "이 문서는 word_parser 테스트용 샘플 문서입니다. "
        "다양한 스타일, 표, 이미지가 포함되어 있습니다.",
        style="Normal",
    )

    # H2
    doc.add_heading("1.1 배경", level=2)
    doc.add_paragraph("파서가 올바르게 청크를 분리하는지 검증하기 위한 섹션입니다.")

    # H3 + table
    doc.add_heading("1.1.1 주요 지표", level=3)
    doc.add_paragraph("아래 표는 핵심 성능 지표를 나타냅니다.")

    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = "Table Grid"
    headers = ["지표", "목표값", "달성값"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    data = [
        ("처리 속도", "100 doc/s", "112 doc/s"),
        ("메모리 사용", "< 512 MB", "380 MB"),
        ("오류율", "< 0.1%", "0.04%"),
    ]
    for r_idx, (a, b, c) in enumerate(data, start=1):
        row = tbl.rows[r_idx].cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    # H3 + image with caption
    doc.add_heading("1.1.2 시스템 구성도", level=3)
    doc.add_paragraph("아래 그림은 전체 시스템 구성을 나타냅니다.")

    png = _png_bytes((70, 130, 180), label="System Architecture")
    doc.add_picture(io.BytesIO(png), width=Inches(3))
    cap = doc.add_paragraph("그림 1. 시스템 구성도", style="Caption")

    # ── H1: 기능 설명 ────────────────────────────────────────────────────────
    doc.add_heading("2. 기능 설명", level=1)
    doc.add_paragraph("각 기능별 상세 설명입니다.")

    # H2
    doc.add_heading("2.1 문서 파싱", level=2)
    doc.add_paragraph(
        "word_parser는 .docx 파일을 읽어 단락, 표, 이미지를 "
        "구조적으로 추출합니다."
    )

    # H3 + image without caption
    doc.add_heading("2.1.1 파싱 흐름", level=3)
    doc.add_paragraph("파싱 흐름을 도식화하면 아래와 같습니다.")

    png2 = _png_bytes((180, 80, 60), label="Parse Flow")
    doc.add_picture(io.BytesIO(png2), width=Inches(3))
    # no caption — tests caption-less image path

    # H3 + table
    doc.add_heading("2.1.2 지원 포맷", level=3)
    doc.add_paragraph("현재 지원하는 이미지 포맷 목록입니다.")

    tbl2 = doc.add_table(rows=5, cols=2)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["포맷", "비고"]):
        tbl2.rows[0].cells[i].text = h
        tbl2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    fmt_data = [("PNG", "기본 지원"), ("JPEG", "기본 지원"), ("SVG", "벡터"), ("EMF", "Windows 전용")]
    for r_idx, (f, n) in enumerate(fmt_data, start=1):
        tbl2.rows[r_idx].cells[0].text = f
        tbl2.rows[r_idx].cells[1].text = n

    # H2
    doc.add_heading("2.2 청크 분리", level=2)
    doc.add_paragraph("헤딩 레벨을 기준으로 문서를 청크 단위로 분리합니다.")

    doc.add_heading("2.2.1 분리 규칙", level=3)
    doc.add_paragraph(
        "split_depth 설정에 따라 해당 레벨의 헤딩마다 새 청크가 생성됩니다. "
        "하위 헤딩은 청크 내 마크다운 헤더로 렌더링됩니다."
    )

    png3 = _png_bytes((60, 160, 90), label="Chunk Split")
    doc.add_picture(io.BytesIO(png3), width=Inches(3))
    doc.add_paragraph("그림 2. 청크 분리 예시", style="Caption")

    # ── H1: 결론 ────────────────────────────────────────────────────────────
    doc.add_heading("3. 결론", level=1)
    doc.add_paragraph("word_parser는 다양한 문서 구조를 안정적으로 처리합니다.")

    doc.add_heading("3.1 향후 계획", level=2)
    doc.add_paragraph("추가 포맷 지원 및 성능 최적화를 예정하고 있습니다.")

    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.style = "Table Grid"
    for i, h in enumerate(["과제", "일정"]):
        tbl3.rows[0].cells[i].text = h
        tbl3.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    plan_data = [("WebP 지원", "2026 Q3"), ("성능 벤치마크", "2026 Q4")]
    for r_idx, (t, d) in enumerate(plan_data, start=1):
        tbl3.rows[r_idx].cells[0].text = t
        tbl3.rows[r_idx].cells[1].text = d

    # ── H1: 병합 셀 예시 ────────────────────────────────────────────────────
    doc.add_heading("4. 병합 셀 예시", level=1)
    doc.add_paragraph("가로/세로 병합 셀이 포함된 표 예시입니다.")

    # 가로 병합: 헤더 행에서 두 컬럼 병합
    doc.add_heading("4.1 가로 병합", level=2)
    tbl_h = doc.add_table(rows=3, cols=3)
    tbl_h.style = "Table Grid"
    # 첫 행: [병합된 헤더(A+B), C]
    tbl_h.rows[0].cells[0].merge(tbl_h.rows[0].cells[1])
    tbl_h.rows[0].cells[0].text = "병합된 헤더"
    tbl_h.rows[0].cells[2].text = "C"
    tbl_h.rows[1].cells[0].text = "A1"
    tbl_h.rows[1].cells[1].text = "B1"
    tbl_h.rows[1].cells[2].text = "C1"
    tbl_h.rows[2].cells[0].text = "A2"
    tbl_h.rows[2].cells[1].text = "B2"
    tbl_h.rows[2].cells[2].text = "C2"

    # 세로 병합: 첫 컬럼 두 행 병합
    doc.add_heading("4.2 세로 병합", level=2)
    tbl_v = doc.add_table(rows=3, cols=2)
    tbl_v.style = "Table Grid"
    tbl_v.rows[0].cells[0].text = "헤더1"
    tbl_v.rows[0].cells[1].text = "헤더2"
    # 세로 병합: rows[1].cells[0] + rows[2].cells[0]
    # 세로 병합: rows[1].cells[0] → restart, rows[2].cells[0] → continuation
    # _tbl.tr_lst 로 raw XML 접근 (python-docx .merge() 는 continuation vMerge 버그 있음)
    tr1_tcs = tbl_v._tbl.tr_lst[1].tc_lst
    tr2_tcs = tbl_v._tbl.tr_lst[2].tc_lst
    _set_vmerge(tr1_tcs[0], restart=True)
    _set_vmerge(tr2_tcs[0], restart=False)
    tbl_v.rows[1].cells[0].text = "병합됨"
    tbl_v.rows[1].cells[1].text = "값1"
    tbl_v.rows[2].cells[1].text = "값2"

    return doc


if __name__ == "__main__":
    out = Path(__file__).parent / "sample.docx"
    doc = build()
    doc.save(out)
    print(f"Saved: {out}")
