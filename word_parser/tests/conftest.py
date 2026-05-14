import io
import pytest
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK


def make_docx() -> Document:
    return Document()


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    doc.add_paragraph(text, style=style)


def add_paragraph_with_font_size(doc: Document, text: str, size_pt: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size_pt)


def add_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(break_type=WD_BREAK.PAGE)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            tbl.rows[r_idx].cells[c_idx].text = cell_text


def docx_bytes(doc: Document) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
